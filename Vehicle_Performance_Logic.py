import math
import io
import os
from datetime import datetime
from typing import Dict, List, Any

# --- Try to import required libraries ---
try:
    import numpy as np
except ImportError:
    print("="*50)
    print("ERROR: 'numpy' library not found for VehiclePerformance.")
    print("Install it using: pip install numpy")
    print("="*50)
    np = None

try:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("="*50)
    print("ERROR: 'python-docx' library not found.")
    print("Install it using: pip install python-docx")
    print("="*50)
    docx = None

# ===============================================================
# Core Physics Functions
# (Extracted from vehicle_performance_calculator.py)
# Units:
# - Speed: km/h
# - Weight: metric tons
# - Forces: Newtons (N)
# - Torque: Newton-meters (Nm), RPM in rev/min, Power in kW
# - Slope: percent (%), curve: degree
# ===============================================================

def rolling_resistance_loco(speed_kmh: float, weight_ton: float, num_axles: int) -> float:
    """Rolling resistance for locomotive in Newtons."""
    if weight_ton <= 0 or num_axles <= 0:
        return 0.0
    A = 0.647 + (13.17 / (weight_ton / num_axles))
    B = 0.00933
    C = 0.057 / weight_ton
    return (A + B * speed_kmh + C * speed_kmh ** 2) * weight_ton * 9.81

def rolling_resistance_wagon(speed_kmh: float, weight_ton: float) -> float:
    """Rolling resistance for wagon in Newtons."""
    if weight_ton <= 0:
        return 0.0
    A = 0.6438797
    B = 0.01047218
    C = 0.00007323
    return (A + B * speed_kmh + C * speed_kmh ** 2) * weight_ton * 9.81

def gradient_resistance(weight_ton: float, slope_percent: float) -> float:
    """Gradient (grade) resistance in Newtons."""
    return weight_ton * 1000 * 9.81 * slope_percent / 100.0

def curvature_resistance(weight_ton: float, curve_degree: float) -> float:
    """Curvature resistance in Newtons."""
    return 0.4 * weight_ton * curve_degree * 9.81

def starting_resistance_loco(weight_ton: float) -> float:
    """Starting resistance for locomotive in Newtons."""
    return 6.0 * weight_ton * 9.81

def starting_resistance_wagon(weight_ton: float) -> float:
    """Starting resistance for wagon in Newtons."""
    return 4.0 * weight_ton * 9.81

def interpolate_torque(engine_rpm: float, curve: Dict[int, float]) -> float:
    """Interpolate engine torque (Nm) at a given RPM from a discrete curve."""
    if not curve:
        raise ValueError("Torque curve is empty.")
    if np is None:
        raise ImportError("Numpy is required for interpolation.")
        
    # Ensure keys are sorted for interpolation
    rpms = np.array(sorted(curve.keys()), dtype=float)
    torques = np.array([curve[r] for r in rpms], dtype=float)
    
    # np.interp handles extrapolation (clamping to endpoints)
    return float(np.interp(engine_rpm, rpms, torques))


# ===============================================================
# Main Calculator Class for FastAPI
# ===============================================================

class VehiclePerformanceCalculator:
    """
    Handles all business logic for Vehicle Performance calculations.
    This class is separate from the FastAPI web server logic.
    """
    
    def __init__(self, inputs: dict):
        """
        Initializes the calculator with validated and processed inputs
        from the main.py validation function.
        """
        self.inputs = inputs
        
        # --- Store key parameters from the validated inputs dict ---
        # Track
        self.max_curve_deg = inputs.get('max_curve', 0.0)
        if inputs.get('curve_unit') == 'm' and self.max_curve_deg != 0:
            self.max_curve_deg = 1750.0 / self.max_curve_deg
            
        self.max_slope_percent = inputs.get('max_slope', 0.0)
        if inputs.get('slope_unit') == 'degree':
            self.max_slope_percent = math.tan(math.radians(self.max_slope_percent)) * 100.0

        # Vehicle
        self.loco_gvw_kg = inputs.get('loco_gvw_kg', 0.0)
        self.loco_gvw_ton = self.loco_gvw_kg / 1000.0
        self.max_speed_kmh = inputs.get('max_speed_kmh', 0.0)
        self.num_axles = inputs.get('num_axles', 1)
        self.rear_axle_ratio = inputs.get('rear_axle_ratio', 1.0)
        self.gear_ratios = inputs.get('gear_ratios', [1.0])
        self.shunting_load_t = inputs.get('shunting_load_t', 0.0)
        
        # RRV
        self.peak_power_kw = inputs.get('peak_power_kw', 0.0)
        self.friction_mu = inputs.get('friction_mu', 0.3)
        self.wheel_dia_m = inputs.get('wheel_dia_m', 0.73)
        self.min_rpm = inputs.get('min_rpm', 100)
        self.max_rpm = inputs.get('max_rpm', 2500)
        
        # Torque Curve (comes in as Dict[int, float])
        self.torque_curve = inputs.get('torque_curve', {})
        if not self.torque_curve:
            raise ValueError("Torque Curve is empty or missing.")
            
        self.max_torque_nm = max(self.torque_curve.values()) if self.torque_curve else 0.0

    def run_tractive_calculation(self) -> dict:
        """
        Computes traction vs slipping limits.
        Adapted from Tkinter 'run_tractive_calculation'.
        """
        if self.wheel_dia_m == 0:
             raise ValueError("Wheel Diameter cannot be zero.")
             
        max_traction_generated_n = 2 * (self.max_torque_nm * max(self.gear_ratios) * self.rear_axle_ratio) / self.wheel_dia_m
        max_traction_slipping_n = self.loco_gvw_ton * self.friction_mu * 1000.0 * 9.81
        
        if max_traction_generated_n > max_traction_slipping_n:
            result_message = "Limited by slipping"
        else:
            result_message = "Not limited by slipping"
            
        return {
            "max_traction_generated_n": max_traction_generated_n,
            "max_traction_slipping_n": max_traction_slipping_n,
            "result_message": result_message
        }

    def _calculate_curves_common(self, calculate_shunting: bool) -> List[dict]:
        """
        Core plotting engine.
        Adapted from Tkinter 'calculate_common'.
        Returns a list of data points for plotting.
        """
        if np is None:
            raise ImportError("Numpy is required to calculate plot data.")
        
        plot_data = []
        
        # Iterate over a range of slopes
        for slope_percent in np.arange(0, self.max_slope_percent + 0.5, 0.5):
            for gear_ratio in self.gear_ratios:
                if self.rear_axle_ratio <= 0 or self.wheel_dia_m <= 0 or gear_ratio <= 0:
                    continue
                
                # 0 to max permitted by RPM
                min_speed_kmh = 0.0
                max_speed_kmh = ((self.max_rpm * math.pi * self.wheel_dia_m) / (gear_ratio * self.rear_axle_ratio * 60.0)) * 3.6
                
                # Generate 100 points for this curve
                speeds_kmh = np.linspace(min_speed_kmh, max_speed_kmh, 100)
                
                for speed_kmh in speeds_kmh:
                    speed_mps = speed_kmh / 3.6
                    wheel_circumference_m = math.pi * self.wheel_dia_m
                    wheel_rpm = (speed_mps / wheel_circumference_m * 60.0) if wheel_circumference_m > 0 else 0
                    engine_rpm = wheel_rpm * gear_ratio * self.rear_axle_ratio

                    # Interpolate torque and limit by power
                    torque_at_rpm = interpolate_torque(engine_rpm, self.torque_curve)
                    power_at_rpm_kw = (engine_rpm * torque_at_rpm * 2.0 * math.pi) / (60.0 * 1000.0)
                    
                    if power_at_rpm_kw > self.peak_power_kw and engine_rpm > 0:
                        torque_at_rpm = (self.peak_power_kw * 60.0 * 1000.0) / (engine_rpm * 2.0 * math.pi)

                    # Calculate Tractive Effort and cap by slipping
                    max_traction_generated_n = 2.0 * (torque_at_rpm * gear_ratio * self.rear_axle_ratio) / self.wheel_dia_m
                    max_traction_slipping_n = self.loco_gvw_ton * self.friction_mu * 1000.0 * 9.81
                    actual_traction_n = min(max_traction_generated_n, max_traction_slipping_n)

                    # Calculate loco resistance (includes starting resistance)
                    loco_resistance = (
                        rolling_resistance_loco(speed_kmh, self.loco_gvw_ton, self.num_axles)
                        + gradient_resistance(self.loco_gvw_ton, slope_percent)
                        + curvature_resistance(self.loco_gvw_ton, self.max_curve_deg)
                        + starting_resistance_loco(self.loco_gvw_ton)
                    )

                    y_val = 0.0
                    if calculate_shunting:
                        remaining_traction_n = actual_traction_n - loco_resistance
                        if remaining_traction_n > 0:
                            # Resistance per 1 ton of wagon
                            total_resistance_1ton_wagon = (
                                rolling_resistance_wagon(speed_kmh, 1)
                                + gradient_resistance(1, slope_percent)
                                + curvature_resistance(1, self.max_curve_deg)
                                + starting_resistance_wagon(1)
                            )
                            if total_resistance_1ton_wagon > 0:
                                y_val = remaining_traction_n / total_resistance_1ton_wagon
                    else:
                        # Tractive effort curve (N)
                        y_val = actual_traction_n

                    plot_data.append({
                        "slope": round(slope_percent, 2),
                        "gear": gear_ratio,
                        "speed_kmh": round(speed_kmh, 2),
                        "value": round(y_val, 2)
                    })
        return plot_data

    def calculate_plot_data(self) -> dict:
        """
        Generates data for both plots expected by the frontend.
        """
        tractive_effort_data = self._calculate_curves_common(calculate_shunting=False)
        shunting_capability_data = self._calculate_curves_common(calculate_shunting=True)
        
        return {
            "tractive_effort_plot": tractive_effort_data,
            "shunting_capability_plot": shunting_capability_data
        }

    def calculate_speed_for_shunting_load(self) -> List[dict]:
        """
        Compute max achievable speed (km/h) vs slope (%) for a given shunting load.
        Adapted from Tkinter 'calculate_speed_for_shunting_load'.
        """
        if np is None:
            raise ImportError("Numpy is required to calculate plot data.")

        table_data = []
        
        # Using the tallest ratio for min speed and shortest for max speed
        if self.rear_axle_ratio <= 0 or self.wheel_dia_m <= 0:
             raise ValueError("Axle Ratio or Wheel Diameter cannot be zero.")

        min_speed_kmh = (self.min_rpm * math.pi * self.wheel_dia_m) / (max(self.gear_ratios) * self.rear_axle_ratio * 60.0) * 3.6
        max_speed_kmh = (self.max_rpm * math.pi * self.wheel_dia_m) / (min(self.gear_ratios) * self.rear_axle_ratio * 60.0) * 3.6
        
        speeds_kmh = np.linspace(min_speed_kmh, max_speed_kmh, 100)

        for slope_percent in np.arange(0, self.max_slope_percent + 0.5, 0.5):
            max_achievable_speed = 0.0

            for speed_kmh in speeds_kmh:
                speed_mps = speed_kmh / 3.6
                wheel_circ = math.pi * self.wheel_dia_m
                wheel_rpm = (speed_mps / wheel_circ * 60.0) if wheel_circ > 0 else 0
                # Using max gear ratio for this calculation (as in Tkinter file)
                engine_rpm = wheel_rpm * max(self.gear_ratios) * self.rear_axle_ratio

                tq = interpolate_torque(engine_rpm, self.torque_curve)

                # Limit by power
                power_kw = (engine_rpm * tq * 2.0 * math.pi) / (60.0 * 1000.0)
                if power_kw > self.peak_power_kw and engine_rpm > 0:
                    tq = (self.peak_power_kw * 60.0 * 1000.0) / (engine_rpm * 2.0 * math.pi)

                # Calculate Tractive Effort and cap by slipping
                max_traction_generated_n = 2 * (tq * max(self.gear_ratios) * self.rear_axle_ratio) / self.wheel_dia_m
                max_traction_slipping_n = self.loco_gvw_ton * self.friction_mu * 1000.0 * 9.81
                actual_traction_n = min(max_traction_generated_n, max_traction_slipping_n)

                # NOTE: This calc uses RUNNING resistance, NOT starting resistance
                # (as per the logic in the Tkinter file 'calculate_speed_for_shunting_load')
                loco_res = (
                    rolling_resistance_loco(speed_kmh, self.loco_gvw_ton, self.num_axles)
                    + gradient_resistance(self.loco_gvw_ton, slope_percent)
                    + curvature_resistance(self.loco_gvw_ton, self.max_curve_deg)
                )

                wagon_res = (
                    rolling_resistance_wagon(speed_kmh, self.shunting_load_t)
                    + gradient_resistance(self.shunting_load_t, slope_percent)
                    + curvature_resistance(self.shunting_load_t, self.max_curve_deg)
                )

                total_resistance_n = loco_res + wagon_res
                
                if actual_traction_n >= total_resistance_n:
                    max_achievable_speed = float(speed_kmh)

            table_data.append({
                "slope": round(slope_percent, 2),
                "max_speed_kmh": round(max_achievable_speed, 2)
            })
        return table_data
        
    def calculate_detailed_step_by_step(self) -> dict:
        """
        Calculate detailed step-by-step calculations with formulas and intermediate results.
        Returns a dictionary with all calculation steps for documentation.
        """
        # Sample calculation at 10 km/h (or minimum speed)
        calc_speed_kmh = 10.0
        
        # Rolling Resistance Calculation
        A = 0.647 + (13.17 / (self.loco_gvw_ton / self.num_axles))
        B = 0.00933
        C = 0.057 / self.loco_gvw_ton
        rolling_res = rolling_resistance_loco(calc_speed_kmh, self.loco_gvw_ton, self.num_axles)
        
        # Gradient Resistance
        gradient_res = gradient_resistance(self.loco_gvw_ton, 0.0)  # At 0% slope
        
        # Curve Resistance
        curve_res = curvature_resistance(self.loco_gvw_ton, self.max_curve_deg)
        
        # Starting Resistance
        starting_res = starting_resistance_loco(self.loco_gvw_ton)
        
        # Total Locomotive Resistance
        total_loco_resistance = rolling_res + gradient_res + curve_res + starting_res
        
        # Max Tractive Effort (Generated)
        max_torque = max(self.torque_curve.values()) if self.torque_curve else 0
        max_gear_ratio = max(self.gear_ratios) if self.gear_ratios else 1
        wheel_radius = self.wheel_dia_m / 2
        max_te_generated = (max_torque * max_gear_ratio * self.rear_axle_ratio) / wheel_radius
        
        # Adhesion Limit
        adhesion_limit = self.loco_gvw_kg * 9.81 * self.friction_mu
        
        # Final Usable Tractive Effort
        final_te = min(max_te_generated, adhesion_limit)
        
        # Net Tractive Effort
        net_te = final_te - total_loco_resistance
        
        # Wagon resistance per ton (at calc speed)
        wagon_res_per_ton = rolling_resistance_wagon(calc_speed_kmh, 1.0) + starting_resistance_wagon(1.0)
        
        # Shunting Capability
        shunting_capability = net_te / wagon_res_per_ton if wagon_res_per_ton > 0 else 0
        
        # Calculate max achievable speed with shunting load
        max_speed_with_load = 0.0
        if self.shunting_load_t > 0:
            for test_speed in range(1, 51):  # Test speeds from 1 to 50 km/h
                test_loco_res = (rolling_resistance_loco(test_speed, self.loco_gvw_ton, self.num_axles) +
                               gradient_res + curve_res)
                test_wagon_res = rolling_resistance_wagon(test_speed, self.shunting_load_t)
                test_total_res = test_loco_res + test_wagon_res
                
                if final_te >= test_total_res:
                    max_speed_with_load = float(test_speed)
        
        return {
            "calculation_speed": calc_speed_kmh,
            "rolling_resistance": {
                "formula": "RR (N) = (A + B×v + C×v²) × W_ton × g",
                "A": round(A, 3),
                "B": B,
                "C": round(C, 5),
                "calculation": f"RR = ({A:.3f} + {B}×{calc_speed_kmh} + {C:.5f}×{calc_speed_kmh}²) × {self.loco_gvw_ton} × 9.81",
                "result": round(rolling_res, 2)
            },
            "gradient_resistance": {
                "formula": "GR (N) = Loco GVW (kg) * g * (Slope / 100)",
                "calculation": f"GR = {self.loco_gvw_kg} × 9.81 × (0.0 ÷ 100)",
                "result": round(gradient_res, 2)
            },
            "curve_resistance": {
                "formula": "CR (N) = 0.4 × W_ton × Curve (deg) × g",
                "calculation": f"CR = 0.4 × {self.loco_gvw_ton} × {self.max_curve_deg} × 9.81",
                "result": round(curve_res, 2)
            },
            "starting_resistance": {
                "formula": "SR (N) = 6.0 × W_ton × g",
                "calculation": f"SR = 6.0 × {self.loco_gvw_ton} × 9.81",
                "result": round(starting_res, 2)
            },
            "total_loco_resistance": {
                "formula": "Total Resistance = RR + GR + CR + Starting Resistance",
                "calculation": f"TR = {rolling_res:.2f} + {gradient_res:.2f} + {curve_res:.2f} + {starting_res:.2f}",
                "result": round(total_loco_resistance, 2)
            },
            "max_tractive_effort": {
                "formula": "TE_max (N) = (Peak Torque × Max Gear Ratio × Axle Ratio) ÷ Wheel Radius",
                "calculation": f"TE = ({max_torque} × {max_gear_ratio} × {self.rear_axle_ratio}) ÷ {wheel_radius:.3f}",
                "result": round(max_te_generated, 2)
            },
            "adhesion_limit": {
                "formula": "Adhesion (N) = Loco GVW (kg) × g × Friction (μ)",
                "calculation": f"Adhesion = {self.loco_gvw_kg} × 9.81 × {self.friction_mu}",
                "result": round(adhesion_limit, 2)
            },
            "final_usable_te": {
                "formula": "Final TE = min(Generated TE, Adhesion Limit)",
                "calculation": f"Final TE = min({max_te_generated:.2f}, {adhesion_limit:.2f})",
                "result": round(final_te, 2)
            },
            "net_tractive_effort": {
                "formula": "Net TE = Final TE - Total Resistance",
                "calculation": f"Net TE = {final_te:.2f} − {total_loco_resistance:.2f}",
                "result": round(net_te, 2)
            },
            "shunting_capability": {
                "formula": "Shuntable Tons = Net TE ÷ Wagon Resistance per Ton",
                "calculation": f"Tons = {net_te:.2f} ÷ {wagon_res_per_ton:.2f}",
                "result": round(shunting_capability, 2)
            },
            "max_achievable_speed": {
                "formula": "Highest speed where Tractive Effort ≥ Total Resistance",
                "load_tons": self.shunting_load_t,
                "result": round(max_speed_with_load, 2)
            }
        }

    def create_report_docx(self):
        """
        Creates a .docx report and returns it as an in-memory stream.
        Adapted from Tkinter 'export_to_docx'.
        """
        if docx is None:
            raise ImportError("python-docx library is required to generate .docx files.")
            
        doc = docx.Document()
        doc.add_heading("Vehicle Performance Calculator Report", level=1)
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_heading("Input Parameters", level=2)
        
        # Use the original validated inputs for the report
        inputs = self.inputs 
        
        # Inputs table
        inputs_table_data = {
            "Loco GVW (kg)": inputs.get('loco_gvw_kg'),
            "Max Speed (km/h)": inputs.get('max_speed_kmh'),
            "Number of Axles": inputs.get('num_axles'),
            "Rear Axle Ratio": inputs.get('rear_axle_ratio'),
            "Gear Ratios": ", ".join(map(str, inputs.get('gear_ratios', []))),
            "Shunting Load (tons)": inputs.get('shunting_load_t'),
            "Peak Power (kW)": inputs.get('peak_power_kw'),
            "Coeff. of Friction (mu)": inputs.get('friction_mu'),
            "Wheel Dia (m)": inputs.get('wheel_dia_m'),
            "Min RPM": inputs.get('min_rpm'),
            "Max RPM": inputs.get('max_rpm'),
            "Max Curve": f"{inputs.get('max_curve')} ({inputs.get('curve_unit')})",
            "Max Slope": f"{inputs.get('max_slope')} ({inputs.get('slope_unit')})",
        }
        
        t_inputs = doc.add_table(rows=1, cols=2)
        t_inputs.style = 'Table Grid'
        t_inputs.rows[0].cells[0].text = "Parameter"
        t_inputs.rows[0].cells[1].text = "Value"
        for label, value in inputs_table_data.items():
            row_cells = t_inputs.add_row().cells
            row_cells[0].text = str(label)
            row_cells[1].text = str(value)
            
        # Torque Curve Table
        doc.add_heading("Torque Curve", level=3)
        t_torque = doc.add_table(rows=1, cols=2)
        t_torque.style = 'Table Grid'
        t_torque.rows[0].cells[0].text = "RPM"
        t_torque.rows[0].cells[1].text = "Torque (Nm)"
        for rpm, torque in sorted(self.torque_curve.items()):
            row_cells = t_torque.add_row().cells
            row_cells[0].text = str(rpm)
            row_cells[1].text = str(torque)

        doc.add_heading("Calculation Results", level=2)
        
        # Add detailed step-by-step calculations
        detailed_calcs = self.calculate_detailed_step_by_step()
        
        doc.add_heading("Functions Used & Step-by-Step Calculations", level=2)
        
        # Rolling Resistance Section
        doc.add_heading(f"Rolling Resistance (at {detailed_calcs['calculation_speed']} km/h)", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['rolling_resistance']['formula']}")
        doc.add_paragraph(f"Calculation: {detailed_calcs['rolling_resistance']['calculation']}")
        doc.add_paragraph(f"Result = {detailed_calcs['rolling_resistance']['result']} N")
        
        # Gradient Resistance Section  
        doc.add_heading("Gradient Resistance", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['gradient_resistance']['formula']}")
        doc.add_paragraph(f"Calculation: {detailed_calcs['gradient_resistance']['calculation']}")
        doc.add_paragraph(f"Result = {detailed_calcs['gradient_resistance']['result']} N")
        
        # Curve Resistance Section
        doc.add_heading("Curve Resistance", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['curve_resistance']['formula']}")
        doc.add_paragraph(f"Calculation: {detailed_calcs['curve_resistance']['calculation']}")
        doc.add_paragraph(f"Result = {detailed_calcs['curve_resistance']['result']} N")
        
        # Starting Resistance Section
        doc.add_heading("Starting Resistance", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['starting_resistance']['formula']}")
        doc.add_paragraph(f"Calculation: {detailed_calcs['starting_resistance']['calculation']}")
        doc.add_paragraph(f"Result = {detailed_calcs['starting_resistance']['result']} N")
        
        # Total Loco Resistance Section
        doc.add_heading(f"Total Loco Resistance (at {detailed_calcs['calculation_speed']} km/h)", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['total_loco_resistance']['formula']}")
        doc.add_paragraph(f"Calculation: {detailed_calcs['total_loco_resistance']['calculation']}")
        doc.add_paragraph(f"Result = {detailed_calcs['total_loco_resistance']['result']} N")
        
        # Max Tractive Effort Section
        doc.add_heading("Max Tractive Effort (Generated)", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['max_tractive_effort']['formula']}")
        doc.add_paragraph(f"Calculation: {detailed_calcs['max_tractive_effort']['calculation']}")
        doc.add_paragraph(f"Result = {detailed_calcs['max_tractive_effort']['result']} N")
        
        # Adhesion Limit Section
        doc.add_heading("Adhesion Limit (Anti-Slip)", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['adhesion_limit']['formula']}")
        doc.add_paragraph(f"Calculation: {detailed_calcs['adhesion_limit']['calculation']}")
        doc.add_paragraph(f"Result = {detailed_calcs['adhesion_limit']['result']} N")
        
        # Final Usable TE Section
        doc.add_heading("Final Usable Tractive Effort", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['final_usable_te']['formula']}")
        doc.add_paragraph(f"Calculation: {detailed_calcs['final_usable_te']['calculation']}")
        doc.add_paragraph(f"Result = {detailed_calcs['final_usable_te']['result']} N")
        
        # Net TE Section
        doc.add_heading(f"Net Tractive Effort (at {detailed_calcs['calculation_speed']} km/h)", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['net_tractive_effort']['formula']}")
        doc.add_paragraph(f"Calculation: {detailed_calcs['net_tractive_effort']['calculation']}")
        doc.add_paragraph(f"Result = {detailed_calcs['net_tractive_effort']['result']} N")
        
        # Shunting Capability Section
        doc.add_heading(f"Shunting Capability (at {detailed_calcs['calculation_speed']} km/h)", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['shunting_capability']['formula']}")
        doc.add_paragraph(f"Calculation: {detailed_calcs['shunting_capability']['calculation']}")
        doc.add_paragraph(f"Result = {detailed_calcs['shunting_capability']['result']} Tons")
        
        # Max Achievable Speed Section
        doc.add_heading(f"Max Achievable Speed (with {detailed_calcs['max_achievable_speed']['load_tons']}t load)", level=3)
        doc.add_paragraph(f"Formula: {detailed_calcs['max_achievable_speed']['formula']}")
        doc.add_paragraph(f"Result = {detailed_calcs['max_achievable_speed']['result']} km/h")
        
        # Final Outputs Summary
        doc.add_heading("Final Outputs", level=2)
        doc.add_paragraph(f"Max Generated Tractive Effort: {detailed_calcs['max_tractive_effort']['result']} N")
        doc.add_paragraph(f"Adhesion Limit (Anti-Slip): {detailed_calcs['adhesion_limit']['result']} N")
        doc.add_paragraph(f"Final Usable Tractive Effort: {detailed_calcs['final_usable_te']['result']} N")
        doc.add_paragraph(f"Total Loco Resistance (at {detailed_calcs['calculation_speed']} km/h): {detailed_calcs['total_loco_resistance']['result']} N")
        doc.add_paragraph(f"Net Tractive Effort (at {detailed_calcs['calculation_speed']} km/h): {detailed_calcs['net_tractive_effort']['result']} N")
        doc.add_paragraph(f"Estimated Shunting Capability (at {detailed_calcs['calculation_speed']} km/h): {detailed_calcs['shunting_capability']['result']} Tons")
        doc.add_paragraph(f"Max Achievable Speed (with load): {detailed_calcs['max_achievable_speed']['result']} km/h")
        
        # Traction Snapshot
        traction_results = self.run_tractive_calculation()
        doc.add_paragraph(f"Max Traction Produced: {traction_results['max_traction_generated_n']:.2f} N")
        doc.add_paragraph(f"Max Traction (No Slip): {traction_results['max_traction_slipping_n']:.2f} N")
        p = doc.add_paragraph()
        run = p.add_run(f"Result: {traction_results['result_message']}")
        run.font.bold = True

        # Speed vs. Slope Table
        doc.add_heading("Max Achievable Speed vs. Slope", level=3)
        doc.add_paragraph(f"(For Shunting Load: {self.shunting_load_t} tons)")
        
        table_data = self.calculate_speed_for_shunting_load()
        t_speed = doc.add_table(rows=1, cols=2)
        t_speed.style = 'Table Grid'
        t_speed.rows[0].cells[0].text = "Slope (%)"
        t_speed.rows[0].cells[1].text = "Max Achievable Speed (km/h)"
        for row in table_data:
            row_cells = t_speed.add_row().cells
            row_cells[0].text = f"{row['slope']:.2f}"
            row_cells[1].text = f"{row['max_speed_kmh']:.2f}"

        # Save to in-memory stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream