import math
import io
from datetime import datetime

try:
    import docx
    from docx.shared import Pt, RGBColor
except ImportError:
    print("="*50)
    print("ERROR: 'python-docx' library not found.")
    print("Install it using: pip install python-docx")
    print("="*50)
    docx = None

# --- Global variables & Constants ---
SIGMA_B_OPTIONS = {
    "880 N/mm²": 880,
    "680 N/mm²": 680,
    "Custom": None
}
CONSTANT_C = 8.257e-7
DEFAULT_V_HEAD = 1.1
KN_TO_TONNES = 1 / 9.80665

class QmaxCalculatorLogic:
    """
    Handles all the business logic for Qmax calculations.
    This class is separate from the FastAPI web server logic.
    """
    def __init__(self):
        # This will store the raw string inputs for use in reports.
        self.inputs_raw = {}

    # --- Core Calculation Logic ---
    def perform_calculations(self, d, sigma_b, v_head):
        """
        Performs the Qmax calculation and returns results.
        """
        qmax_kn = CONSTANT_C * (d / 2) * (sigma_b / v_head) ** 2
        qmax_tonnes = qmax_kn * KN_TO_TONNES
        
        return {
            "d": d,
            "sigma_b": sigma_b,
            "v_head": v_head,
            "qmax_kn": qmax_kn,
            "qmax_tonnes": qmax_tonnes
        }

    # --- Text Formatting for GUI and Report ---
    def format_detailed_steps(self, results):
        """Formats the step-by-step calculation into a readable string."""
        sigma_v_head_squared = (results['sigma_b'] / results['v_head']) ** 2
        d_half = results['d'] / 2
        
        # Get raw inputs for the report
        raw = self.inputs_raw
        
        report_lines = [
            "# Qmax Calculation Report",
            "\n--- INPUT PARAMETERS ---",
            f"Worn rail diameter limit (d): {raw.get('d', 'N/A')} mm",
            f"Material Strength (σB): {raw.get('sigma_b_selection', 'N/A')}",
            f"  (Value Used: {results['sigma_b']} N/mm²)",
            f"Safety Factor (v_head): {raw.get('v_head', 'N/A')}",
            
            "\n--- STEP-BY-STEP CALCULATION ---",
            "1. Formula:",
            "   Qmax = C × (d / 2) × (σB / v_head)²",
            f"   Where C = {CONSTANT_C}",
            
            "\n2. Substitute Values:",
            f"   d = {results['d']} mm",
            f"   σB = {results['sigma_b']} N/mm²",
            f"   v_head = {results['v_head']}",
            
            "\n3. Step-by-Step Calculation:",
            f"   a) (σB / v_head)² = ({results['sigma_b']} / {results['v_head']})² = {sigma_v_head_squared:.3f}",
            f"   b) Qmax = {CONSTANT_C} × ({results['d']} / 2) × {sigma_v_head_squared:.3f}",
            f"   c) Qmax = {CONSTANT_C} × {d_half:.1f} × {sigma_v_head_squared:.3f}",
            
            "\n--- FINAL RESULT ---",
            f"Qmax = {results['qmax_kn']:.4f} kN",
            f"Qmax = {results['qmax_tonnes']:.4f} tonnes"
        ]
        
        return "\n".join(report_lines)

    # --- Docx Report Generation ---
    def create_report_docx(self, results):
        """Creates a .docx report and returns it as an in-memory stream."""
        if docx is None:
            raise ImportError("python-docx library is required to generate .docx files.")
            
        doc = docx.Document()
        doc.add_heading('Qmax Calculation Report', 0)
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.add_heading('1. Input Parameters', level=1)
        raw = self.inputs_raw
        p = doc.add_paragraph()
        p.add_run(f"  • Worn rail diameter limit (d): {raw.get('d')} mm\n")
        p.add_run(f"  • Material Strength (σB): {raw.get('sigma_b_selection')}\n")
        p.add_run(f"    (Value Used: {results['sigma_b']} N/mm²)\n")
        p.add_run(f"  • Safety Factor (v_head): {raw.get('v_head')}\n")

        doc.add_heading('2. Calculation', level=1)
        # We need to format the steps without the input summary for the docx
        sigma_v_head_squared = (results['sigma_b'] / results['v_head']) ** 2
        d_half = results['d'] / 2
        
        detailed_steps_text = (
            f"1. Formula:\n"
            f"   Qmax = C × (d / 2) × (σB / v_head)²\n"
            f"   Where C = {CONSTANT_C}\n\n"
            f"2. Substitute Values:\n"
            f"   d = {results['d']} mm\n"
            f"   σB = {results['sigma_b']} N/mm²\n"
            f"   v_head = {results['v_head']}\n\n"
            f"3. Step-by-Step Calculation:\n"
            f"   a) (σB / v_head)² = ({results['sigma_b']} / {results['v_head']})² = {sigma_v_head_squared:.3f}\n\n"
            f"   b) Qmax = {CONSTANT_C} × ({results['d']} / 2) × {sigma_v_head_squared:.3f}\n"
            f"   c) Qmax = {CONSTANT_C} × {d_half:.1f} × {sigma_v_head_squared:.3f}\n"
        )
        doc.add_paragraph(detailed_steps_text)

        doc.add_heading('3. Final Result', level=1)
        p_res = doc.add_paragraph()
        run_kn = p_res.add_run(f"Qmax = {results['qmax_kn']:.4f} kN\n")
        run_kn.font.bold = True
        run_kn.font.color.rgb = RGBColor.from_string("0D47A1")
        
        p_res.add_run(f"Qmax = {results['qmax_tonnes']:.4f} tonnes")

        # Save to in-memory stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream