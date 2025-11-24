import math
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Global variables ---
# Assume the script is run from where main.py is, so CURRENT_DIR is that directory.
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGE_PATH = os.path.join(CURRENT_DIR, "Diagram.png")

# --- Core Calculation Logic ---
def perform_calculations(config_type, total_load, front_percent, q1_percent, q3_percent):
    """
    Performs all load distribution calculations and returns the results as a dictionary.
    (Copied from your Load Distribution Safety Report.py)
    """
    results = {}
    is_bogie = config_type == "Bogie"
    rear_percent = 100 - front_percent

    front_load = (front_percent / 100) * total_load
    rear_load = total_load - front_load

    q1_val = (q1_percent / 100) * front_load
    q2_val = front_load - q1_val
    q3_val = (q3_percent / 100) * rear_load
    q4_val = rear_load - q3_val

    q_values = {"Q1": q1_val, "Q2": q2_val, "Q3": q3_val, "Q4": q4_val}
    
    results['q_values'] = q_values
    results['front_load'] = front_load
    results['rear_load'] = rear_load
    
    # Find lowest and highest Q values
    results['ql_name'] = min(q_values, key=q_values.get)
    results['ql_value'] = q_values[results['ql_name']]
    results['qh_name'] = max(q_values, key=q_values.get)
    results['qh_value'] = q_values[results['qh_name']]

    # Calculate Q (Average on heavier axle)
    if front_load >= rear_load:
        results['q_formula_str'] = "(Q1 + Q2) / 2"
        results['q_value'] = (q1_val + q2_val) / 2
    else:
        results['q_formula_str'] = "(Q3 + Q4) / 2"
        results['q_value'] = (q3_val + q4_val) / 2

    results['delta_q'] = results['q_value'] - results['ql_value']
    results['delta_q_by_q'] = results['delta_q'] / results['q_value'] if results['q_value'] != 0 else 0
    results['limit'] = 0.6 if is_bogie else 0.5
    
    if results['delta_q_by_q'] <= results['limit']:
        results['status'] = "success"
        results['status_msg'] = f"PASS: ΔQ/Q ({results['delta_q_by_q']:.2%}) is within the {results['limit']:.0%} limit."
    else:
        results['status'] = "fail"
        results['status_msg'] = f"FAIL: ΔQ/Q ({results['delta_q_by_q']:.2%}) exceeds the {results['limit']:.0%} limit."
        
    return results

# --- Text Formatting for GUI and Report ---
def format_detailed_steps(inputs, results):
    """
    Formats the step-by-step calculation into a readable string.
    (Copied from your Load Distribution Safety Report.py)
    """
    return (
        f"1. Calculate Front and Rear Loads:\n"
        f"   Front = Total Load × (Front % / 100)\n"
        f"   Front = {inputs['total_load']:.2f} × ({inputs['front_percent']:.2f} / 100) = {results['front_load']:.2f} Ton\n\n"
        f"   Rear = Total Load - Front Load\n"
        f"   Rear = {inputs['total_load']:.2f} - {results['front_load']:.2f} = {results['rear_load']:.2f} Ton\n\n"
        
        f"2. Calculate Individual Wheel Loads (Q Values):\n"
        f"   Q1 = Front Load × (Q1 % / 100)\n"
        f"   Q1 = {results['front_load']:.2f} × ({inputs['q1_percent']:.2f} / 100) = {results['q_values']['Q1']:.2f} Ton\n\n"
        f"   Q2 = Front Load - Q1\n"
        f"   Q2 = {results['front_load']:.2f} - {results['q_values']['Q1']:.2f} = {results['q_values']['Q2']:.2f} Ton\n\n"
        f"   Q3 = Rear Load × (Q3 % / 100)\n"
        f"   Q3 = {results['rear_load']:.2f} × ({inputs['q3_percent']:.2f} / 100) = {results['q_values']['Q3']:.2f} Ton\n\n"
        f"   Q4 = Rear Load - Q3\n"
        f"   Q4 = {results['rear_load']:.2f} - {results['q_values']['Q3']:.2f} = {results['q_values']['Q4']:.2f} Ton\n\n"
        
        f"3. Calculate Q, ΔQ, and ΔQ/Q:\n"
        f"   Q (Average on heavier axle) = {results['q_value']:.2f} Ton\n"
        f"   QL (Lowest wheel load) = {results['ql_value']:.2f} Ton ({results['ql_name']})\n"
        f"   ΔQ = Q - QL = {results['q_value']:.2f} - {results['ql_value']:.2f} = {results['delta_q']:.2f} Ton\n\n"
        f"   ΔQ/Q = ΔQ / Q = {results['delta_q']:.2f} / {results['q_value']:.2f} = {results['delta_q_by_q']:.4f}\n\n"
        
        f"4. Final Check:\n"
        f"   Is {results['delta_q_by_q']:.2%} ≤ {results['limit']:.0%}? {'Yes' if results['status']=='success' else 'No'}.\n"
        f"   Result: {results['status'].upper()}"
    )

# --- Docx Report Generation ---
def create_report_docx(inputs, results):
    """
    Creates a .docx report and returns it as an in-memory stream.
    (Adapted from your Load Distribution Safety Report.py)
    """
    doc = Document()
    doc.add_heading('Load Distribution Safety Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('1. Input Parameters', level=1)
    p = doc.add_paragraph()
    p.add_run(f"  • Configuration Type: {inputs['config_type']}\n")
    p.add_run(f"  • Total Load: {inputs['total_load']:.2f} Ton\n")
    p.add_run(f"  • Front Load Percentage: {inputs['front_percent']:.2f}%\n")
    p.add_run(f"  • Q1 Percentage (of Front Load): {inputs['q1_percent']:.2f}%\n")
    p.add_run(f"  • Q3 Percentage (of Rear Load): {inputs['q3_percent']:.2f}%\n")

    doc.add_heading('2. Calculation Results Summary', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(4.0)
    table.columns[1].width = Inches(2.5)
    cell_text = (
        f"Overall Status: {results['status'].upper()}\n\n"
        f"ΔQ/Q Ratio: {results['delta_q_by_q']:.2%}\n"
        f"Allowed Limit: {results['limit']:.0%}\n\n"
        f"Front Load: {results['front_load']:.2f} Ton\n"
        f"Rear Load: {results['rear_load']:.2f} Ton\n\n"
        f"Q1: {results['q_values']['Q1']:.2f} Ton | Q2: {results['q_values']['Q2']:.2f} Ton\n"
        f"Q3: {results['q_values']['Q3']:.2f} Ton | Q4: {results['q_values']['Q4']:.2f} Ton"
    )
    table.cell(0, 0).text = cell_text
    p = table.cell(0, 1).paragraphs[0]
    
    # Check for image and add it if it exists
    if os.path.exists(IMAGE_PATH):
        try:
            p.add_run().add_picture(IMAGE_PATH, width=Inches(2.5))
        except Exception as e:
            p.text = f"Diagram.png found but could not be added. Error: {e}"
    else:
        p.text = "Diagram.png not found."

    doc.add_heading('3. Detailed Calculation Steps', level=1)
    detailed_steps_text = format_detailed_steps(inputs, results)
    doc.add_paragraph(detailed_steps_text)

    try:
        # Save to in-memory stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        # Raise an exception that can be caught by FastAPI
        raise RuntimeError(f"An error occurred while creating the report: {e}")