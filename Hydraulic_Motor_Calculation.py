import math
import io
try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("="*50)
    print("ERROR: 'python-docx' library not found.")
    print("Install it using: pip install python-docx")
    print("="*50)
    # If this script is imported, we might not want to exit,
    # but the docx functions will fail.
    # We'll let it fail later if docx is None.
    docx = None

class Calculator:
    """
    Handles all the business logic for hydraulic motor calculations.
    This class is now separate from the FastAPI web server logic.
    """
    def __init__(self):
        # This will store the raw string inputs for use in reports.
        self.inputs_raw = {}

    # --- CALCULATION MODE 1 ---
    def perform_displacement_calculation(self, inputs: dict) -> dict:
        results = {}
        pump_results = []
        warnings = [] 
        speed = inputs['speed']
        locomotive_weight = inputs['weight']
        number_of_axles = inputs['axles']
        wheel_diameter = inputs['wheel_diameter']
        slope_percent = inputs['slope_percent']
        curve_degree = inputs['curve_degree']
        pressure_bar = inputs['pressure']
        mechanical_efficiency = inputs['mech_eff_motor'] / 100.0 
        gear_ratio = inputs.get('axle_gear_box_ratio', 1.0)
        engine_gear_list = inputs.get('engine_gear_ratio_list', [1.0])
        pto_gear_ratio = inputs['pto_gear_ratio']
        gravity = 9.81
        wheel_circumference = (wheel_diameter * math.pi) / 1000
        if wheel_circumference == 0: raise ValueError("Wheel Diameter cannot be 0.")
        speed_mps = (speed * 1000) / 3600
        wheel_rpm = (speed_mps / wheel_circumference) * 60
        gearbox_input_rpm = wheel_rpm * gear_ratio
        if number_of_axles <= 0: raise ValueError("Number of axles must be greater than 0.")
        if locomotive_weight <= 0: raise ValueError("Weight must be greater than 0.")
        A = 0.647 + (13.17 / (locomotive_weight / number_of_axles))
        B = 0.00933
        C = 0.057 / locomotive_weight
        rolling_resistance = (A + B * speed + C * speed**2) * locomotive_weight * gravity / 1000
        gradient_resistance = locomotive_weight * 1000 * gravity * slope_percent / 100000
        curvature_resistance = 0.4 * locomotive_weight * curve_degree * gravity / 1000
        starting_resistance = 6 * locomotive_weight * gravity / 1000
        total_resistance = rolling_resistance + gradient_resistance + curvature_resistance + starting_resistance
        wheel_radius = wheel_diameter / 2000
        required_total_torque = total_resistance * 1000 * wheel_radius
        number_of_wheels = number_of_axles * 2
        per_wheel_torque = required_total_torque / number_of_wheels if number_of_wheels > 0 else 0.0
        per_axle_torque = required_total_torque / number_of_axles if number_of_axles > 0 else 0.0
        if gear_ratio == 0: raise ValueError("Axle Gearbox Ratio cannot be 0.")
        per_gearbox_input_torque = per_axle_torque / gear_ratio
        per_gearbox_input_torque_kg_cm = per_gearbox_input_torque * 10.1972
        pressure_kg_cm2 = pressure_bar * 1.01972
        if pressure_kg_cm2 == 0 or mechanical_efficiency == 0:
            raise ValueError("Pressure or Mechanical Efficiency cannot be 0.")
        motor_displacement = (per_gearbox_input_torque_kg_cm * 2 * 3.1416) / (pressure_kg_cm2 * mechanical_efficiency)
        vol_eff_motor_frac = (inputs['vol_eff_motor'] / 100.0)
        if vol_eff_motor_frac == 0: raise ValueError("Motor Volumetric Efficiency cannot be 0.")
        hydraulic_motor_flow = ((motor_displacement * gearbox_input_rpm) / vol_eff_motor_frac ) / 1000
        results['speed_mps'] = speed_mps
        results['wheel_circumference'] = wheel_circumference
        results['wheel_rpm'] = wheel_rpm
        results['gearbox_input_rpm'] = gearbox_input_rpm
        results['A'] = A; results['B'] = B; results['C'] = C
        results['rolling_resistance'] = rolling_resistance
        results['gradient_resistance'] = gradient_resistance
        results['curvature_resistance'] = curvature_resistance
        results['starting_resistance'] = starting_resistance
        results['total_resistance'] = total_resistance
        results['wheel_radius'] = wheel_radius
        results['required_total_torque'] = required_total_torque
        results['per_wheel_torque'] = per_wheel_torque
        results['per_axle_torque'] = per_axle_torque
        results['per_gearbox_input_torque'] = per_gearbox_input_torque
        results['per_gearbox_input_torque_kg_cm'] = per_gearbox_input_torque_kg_cm
        results['pressure_kg_cm2'] = pressure_kg_cm2
        results['motor_displacement_cc'] = motor_displacement
        results['per_motor_flow_rate_lpm'] = hydraulic_motor_flow
        prop_rpm_list = [inputs.get('max_vehicle_rpm', 0.0)]
        vol_eff_pump_frac = inputs['vol_eff_pump'] / 100.0
        if vol_eff_pump_frac == 0: raise ValueError("Pump Volumetric Efficiency cannot be 0.")
        for engine_gear in engine_gear_list:
            for max_vehicle_rpm_input in prop_rpm_list:
                if engine_gear == 0 or gear_ratio == 0:
                        raise ValueError("Engine or Axle Gear Ratio cannot be 0.")
                actual_prop_rpm = max_vehicle_rpm_input / engine_gear
                pump_rpm_from_prop = pto_gear_ratio * actual_prop_rpm
                if pump_rpm_from_prop <= 0:
                    raise ValueError(f"Calculated pump RPM ({pump_rpm_from_prop:.2f}) is 0 or negative.")
                pump_denom = (pump_rpm_from_prop * vol_eff_pump_frac)
                if pump_denom == 0:
                        raise ValueError("Calculated Pump RPM or Volumetric Efficiency is 0.")
                disp_L_rev = hydraulic_motor_flow / pump_denom
                pump_disp = disp_L_rev * 1000.0
                pump_results.append({
                    'engine_gear_ratio': engine_gear, 'max_vehicle_rpm_input': max_vehicle_rpm_input, 
                    'pump_rpm': pump_rpm_from_prop, 'pump_disp_cc': pump_disp, 'actual_prop_rpm': actual_prop_rpm
                })
        results['pump_results'] = pump_results
        results['warnings'] = warnings
        return results

    # --- CALCULATION MODE 2 ---
    def perform_speed_calculation(self, inputs: dict) -> dict:
        results = {}
        warnings = []
        axle_gear_box_ratio = inputs.get('axle_gear_box_ratio', 1.0)
        engine_gear_list = inputs.get('engine_gear_ratio_list', [1.0])
        if axle_gear_box_ratio <= 0: raise ValueError("Axle Gearbox Ratio must be greater than 0.")
        prop_rpm_list = [inputs.get('max_vehicle_rpm', 0.0)]
        results_list = []
        vol_eff_pump_frac = (inputs['vol_eff_pump'] / 100.0)
        vol_eff_motor_frac = (inputs['vol_eff_motor'] / 100.0)
        if vol_eff_pump_frac == 0: raise ValueError("Pump Volumetric Efficiency cannot be 0.")
        if vol_eff_motor_frac == 0: raise ValueError("Motor Volumetric Efficiency cannot be 0.")
        for engine_gear in engine_gear_list:
            for max_vehicle_rpm_input in prop_rpm_list:
                if engine_gear == 0 or axle_gear_box_ratio == 0:
                        raise ValueError("Engine or Axle Gear Ratio cannot be 0.")
                actual_prop_rpm = max_vehicle_rpm_input / engine_gear
                pump_rpm = actual_prop_rpm * inputs['pto_gear_ratio']
                if pump_rpm > inputs['max_pump_rpm']:
                    warnings.append(f"Engine Gear {engine_gear}: Pump speed ({pump_rpm:.0f} RPM) exceeds max Pump RPM ({inputs['max_pump_rpm']:.0f} RPM).")
                pump_flow_lpm = (inputs['pump_disp_in'] * pump_rpm * vol_eff_pump_frac) / 1000.0
                motor_disp_lpm = inputs['motor_disp_in'] / 1000.0
                if motor_disp_lpm <= 0.0: raise ValueError("Motor displacement cannot be 0 for speed calculation.")
                motor_speed_rpm = (pump_flow_lpm / motor_disp_lpm) * vol_eff_motor_frac
                if motor_speed_rpm > inputs['max_motor_rpm']:
                    warnings.append(f"Engine Gear {engine_gear}: Motor speed ({motor_speed_rpm:.0f} RPM) exceeds max Motor RPM ({inputs['max_motor_rpm']:.0f} RPM).")
                axle_shaft_rpm = motor_speed_rpm / axle_gear_box_ratio
                wheel_circumference = (inputs['wheel_diameter'] * math.pi) / 1000.0
                if wheel_circumference <= 0: raise ValueError("Wheel Diameter must be greater than 0.")
                speed_mps = (axle_shaft_rpm * wheel_circumference) / 60.0
                achievable_speed_kph = speed_mps * 3.6
                wheel_rpm = axle_shaft_rpm
                loop_res = {
                    'engine_gear_ratio': engine_gear, 'max_vehicle_rpm_input': max_vehicle_rpm_input,
                    'actual_prop_rpm': actual_prop_rpm, 'pump_rpm': pump_rpm,
                    'pump_flow_lpm': pump_flow_lpm, 'motor_disp_lpm': motor_disp_lpm,
                    'motor_speed_rpm': motor_speed_rpm, 'axle_shaft_rpm': axle_shaft_rpm,
                    'wheel_circumference': wheel_circumference, 'wheel_rpm': wheel_rpm,
                    'achievable_speed_kph': achievable_speed_kph
                }
                results_list.append(loop_res)
        results['speed_results_list'] = results_list
        results['warnings'] = warnings
        if results_list: results.update(results_list[-1])
        return results

    # --- Report Generator 1 ---
    def _generate_mode1_report(self, inputs: dict, results: dict) -> str:
        raw = self.inputs_raw
        output_lines = []
        output_lines.append("# Pump & Motor (cc) Calculation as per Vehicle")
        output_lines.append("\n--- VEHICLE INPUTS ---")
        output_lines.append(f"Vehicle Weight: {raw.get('weight')} t")
        output_lines.append(f"Number of axles: {raw.get('axles')}")
        output_lines.append(f"Target Speed: {raw.get('speed')} km/h")
        output_lines.append(f"Wheel Dia: {raw.get('wheel_diameter')} mm")
        output_lines.append(f"Slope: {raw.get('slope_percent')} %")
        output_lines.append(f"Curve: {raw.get('curve_degree')} deg")
        output_lines.append(f"Axle Gear box Ratio: {raw.get('axle_gear_box_ratio')}")
        output_lines.append(f"max Vehicle RPM : {raw.get('max_vehicle_rpm')}")
        output_lines.append(f"PTO Gear Box Ratio: {raw.get('pto_gear_ratio')}")
        output_lines.append(f"Engine Gear Box Ratios: {raw.get('engine_gear_ratio')}")
        output_lines.append("\n--- HYDRAULIC MOTOR & PUMP INPUTS ---")
        output_lines.append(f"Total Hydraulic Motor: {raw.get('num_motors')}")
        output_lines.append(f"Hydraulic Motor / axle: {raw.get('per_axle_motor')}")
        output_lines.append(f"Pressure: {raw.get('pressure')} bar")
        output_lines.append(f"Motor Mechanical Efficiency: {raw.get('mech_eff_motor')} %")
        output_lines.append(f"Motor Volumetric Efficiency: {raw.get('vol_eff_motor')} %")
        output_lines.append(f"Pump Volumetric Efficiency: {raw.get('vol_eff_pump')} %")
        output_lines.append("\n--- RESULTS: STEP-BY-STEP CALCULATION (COMMON) ---")
        output_lines.append("\nStep 1: Vehicle Speed & Wheel RPM")
        output_lines.append(f"  Speed (m/s) = {results.get('speed_mps'):.2f} m/s")
        output_lines.append(f"  Wheel Circumference (m) = {results.get('wheel_circumference'):.3f} m")
        output_lines.append(f"  Wheel RPM = {results.get('wheel_rpm'):.2f} RPM")
        output_lines.append("\nStep 2: Resistance Forces (kN)")
        output_lines.append(f"  Rolling Resistance (kN) = {results.get('rolling_resistance'):.2f} kN")
        output_lines.append(f"  Gradient Resistance (kN) = {results.get('gradient_resistance'):.2f} kN")
        output_lines.append(f"  Curvature Resistance (kN) = {results.get('curvature_resistance'):.2f} kN")
        output_lines.append(f"  Starting Resistance (kN) = {results.get('starting_resistance'):.2f} kN")
        output_lines.append("  ---")
        output_lines.append(f"  Total Resistance (kN): {results.get('total_resistance'):.2f} kN")
        output_lines.append("  ---")
        output_lines.append("\nStep 3: Torque Requirements")
        output_lines.append(f"  Wheel Radius (m) = {results.get('wheel_radius'):.3f} m")
        output_lines.append(f"  Required Total Torque (Nm) = {results.get('required_total_torque'):.2f} Nm")
        output_lines.append(f"  Required Torque per Wheel (Nm) = {results.get('per_wheel_torque'):.2f} Nm/wheel")
        output_lines.append(f"  Required Torque per Axle (Nm) = {results.get('per_axle_torque'):.2f} Nm/axle")
        output_lines.append(f"  Required Motor Torque (Nm) = {results.get('per_gearbox_input_torque'):.2f} Nm")
        output_lines.append("\nStep 4: Motor Speed")
        output_lines.append("  Motor Speed (RPM) (gearbox input RPM):")
        output_lines.append(f"    = {results.get('wheel_rpm'):.2f} * {inputs.get('axle_gear_box_ratio')} = {results.get('gearbox_input_rpm'):.2f} RPM")
        output_lines.append("\nStep 5: Motor Displacement (New Formula)")
        output_lines.append(f"  Motor Torque (kg-cm) = {results.get('per_gearbox_input_torque_kg_cm'):.2f} kg-cm")
        output_lines.append(f"  Pressure (kg/cm2) = {results.get('pressure_kg_cm2'):.2f} kg/cm2")
        output_lines.append(f"  Motor Displacement (cc/rev) = {results.get('motor_displacement_cc'):.2f} cc/rev")
        output_lines.append("\nStep 6: Motor Flow Rate (New Formula)")
        output_lines.append(f"  Per Motor Flow Rate (LPM) = {results.get('per_motor_flow_rate_lpm'):.2f} LPM")
        output_lines.append("\n--- RESULTS: STEP-BY-STEP CALCULATION (PER GEAR) ---")
        output_lines.append("\nStep 7: Required Pump Displacement")
        pump_results_list = results.get('pump_results', [])
        if not pump_results_list:
            output_lines.append("    (No pump results were calculated)")
        for res in pump_results_list:
            max_vehicle_rpm_input = res.get('max_vehicle_rpm_input', 0)
            actual_prop_rpm = res.get('actual_prop_rpm', 0)
            calc_pump_rpm = res.get('pump_rpm', 0)
            final_disp_cc = res.get('pump_disp_cc', 0)
            engine_gear = res.get('engine_gear_ratio', 1.0)
            output_lines.append(f"\n  --- For Engine Gear {engine_gear:.2f} @ {max_vehicle_rpm_input:.0f} RPM ---")
            output_lines.append(f"    Actual Prop RPM = {actual_prop_rpm:.2f} RPM")
            output_lines.append(f"    Calculate Pump RPM = {calc_pump_rpm:.2f} RPM")
            output_lines.append(f"    Required Pump Displacement (cc/rev) = {final_disp_cc:.2f} cc/rev")
        return '\n'.join(output_lines)
        
    # --- Report Generator 2 ---
    def _generate_mode2_report(self, inputs: dict, results: dict) -> str:
        raw = self.inputs_raw
        output_lines = []
        output_lines.append("# MODE 2 REPORT: SPEED CALCULATION")
        output_lines.append("\n--- VEHICLE INPUTS ---")
        output_lines.append(f"Wheel Dia: {raw.get('wheel_diameter')} mm")
        output_lines.append(f"Axle Gear box Ratio: {raw.get('axle_gear_box_ratio')}")
        output_lines.append(f"max Vehicle RPM: {raw.get('max_vehicle_rpm')}")
        output_lines.append(f"PTO Gear Box Ratio: {raw.get('pto_gear_ratio')}")
        output_lines.append(f"Engine Gear Box Ratios: {raw.get('engine_gear_ratio')}")
        output_lines.append("\n--- HYDRAULIC MOTOR & PUMP INPUTS ---")
        output_lines.append(f"Total Hydraulic Motor: {raw.get('num_motors')}")
        output_lines.append(f"Hydraulic Motor / axle: {raw.get('per_axle_motor')}")
        output_lines.append(f"Motor Volumetric Efficiency: {raw.get('vol_eff_motor')} %")
        output_lines.append(f"max Motor RPM: {raw.get('max_motor_rpm')} RPM")
        output_lines.append(f"Motor Displacement: {raw.get('motor_disp_in')} cc")
        output_lines.append(f"Pump Volumetric Efficiency: {raw.get('vol_eff_pump')} %")
        output_lines.append(f"max Limit of Pump RPM: {raw.get('max_pump_rpm')} RPM")
        output_lines.append(f"Pump Displacement: {raw.get('pump_disp_in')} cc")
        output_lines.append("\n--- CALCULATION RESULTS (PER ENGINE GEAR) ---")
        speed_results_list = results.get('speed_results_list', [])
        if not speed_results_list:
                 output_lines.append("    (No speed results were calculated)")
        for res in speed_results_list:
            max_vehicle_rpm_input = res.get('max_vehicle_rpm_input', 0)
            actual_prop_rpm = res.get('actual_prop_rpm', 0)
            pump_rpm = res.get('pump_rpm', 0)
            pump_flow_lpm = res.get('pump_flow_lpm', 0)
            motor_speed_rpm = res.get('motor_speed_rpm', 0)
            axle_shaft_rpm = res.get('axle_shaft_rpm', 0)
            achievable_speed_kph = res.get('achievable_speed_kph', 0)
            engine_gear = res.get('engine_gear_ratio', 1.0)
            output_lines.append(f"\n  --- For Engine Gear {engine_gear:.2f} @ {max_vehicle_rpm_input:.0f} RPM ---")
            output_lines.append(f"  Actual Prop RPM = {actual_prop_rpm:.2f} RPM")
            output_lines.append(f"  Calculate Pump RPM = {pump_rpm:.2f} RPM")
            output_lines.append(f"  Calculate Pump Flow (LPM) = {pump_flow_lpm:.2f} LPM")
            output_lines.append(f"  Calculate Motor Speed (RPM) = {motor_speed_rpm:.2f} RPM")
            output_lines.append(f"  Calculate Axle/Wheel Speed (RPM) = {axle_shaft_rpm:.2f} RPM")
            output_lines.append(f"  ** Achievable Speed: {achievable_speed_kph:.2f} km/h **")
        if results.get('warnings'):
            output_lines.append("\n--- ⚠️ WARNINGS ---")
            for warning in results['warnings']:
                output_lines.append(f"- {warning}")
        return '\n'.join(output_lines)

    # --- .docx Creator Helpers ---
    def _add_colored_run(self, paragraph, text, color_rgb, bold=False, size_pt=None):
        if docx is None: return
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor.from_string(color_rgb)
        run.font.bold = bold
        if size_pt:
            run.font.size = Pt(size_pt)
        return run

    def _add_highlighted_result(self, paragraph, label, value, unit):
        if docx is None: return
        paragraph.add_run(label)
        run = paragraph.add_run(f"{value:.2f}{unit}")
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("0D47A1") # Dark Blue
        
    def _populate_input_table(self, doc, title, data_dict):
        if docx is None: return
        doc.add_heading(title, level=3)
        table = doc.add_table(rows=len(data_dict), cols=2)
        table.style = 'Table Grid'
        for i, (key, value) in enumerate(data_dict.items()):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
            table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        return table

    # --- .docx Mode 1 ---
    def _create_mode1_docx(self, doc, inputs, results_wrapper):
        if docx is None: raise ImportError("python-docx library is required to generate .docx files.")
        doc.add_heading("Pump & Motor (cc) Calculation Report", level=0)
        raw = self.inputs_raw
        results = results_wrapper
        
        doc.add_heading("Input Parameters", level=2)
        vehicle_data = {
            "Vehicle Weight (t)": raw.get('weight'), "Number of axles": raw.get('axles'),
            "Target Speed (km/h)": raw.get('speed'), "Slope (%)": raw.get('slope_percent'),
            "Curve (deg)": raw.get('curve_degree'), "Wheel Dia (mm)": raw.get('wheel_diameter'),
            "PTO Gear box Ratio": raw.get('pto_gear_ratio'), "Axle Gear box Ratio": raw.get('axle_gear_box_ratio'),
            "Engine Gear Box Ratios": raw.get('engine_gear_ratio'), "max Vehicle RPM": raw.get('max_vehicle_rpm'),
        }
        hydraulic_data = {
            "Hydraulic Motor (Total)": raw.get('num_motors'), "Hydraulic Motor / axle": raw.get('per_axle_motor'),
            "Pressure (bar)": raw.get('pressure'), "Motor Mech Eff (%)": raw.get('mech_eff_motor'),
            "Motor Vol Eff (%)": raw.get('vol_eff_motor'), "Pump Vol Eff (%)": raw.get('vol_eff_pump'),
        }
        self._populate_input_table(doc, "Vehicle Inputs", vehicle_data)
        self._populate_input_table(doc, "Hydraulic Inputs", hydraulic_data)
        doc.add_page_break()
        
        doc.add_heading("Step-by-Step Calculation (Common)", level=2)
        doc.add_heading("Step 1: Vehicle Speed & Wheel RPM", level=3)
        p = doc.add_paragraph("  Convert Speed (km/h) to Speed (m/s)\n"); p.add_run(f"    Speed (m/s) = ({inputs.get('speed')} * 1000) / 3600\n"); self._add_highlighted_result(p, "    = ", results.get('speed_mps'), " m/s")
        p = doc.add_paragraph("  Calculate Wheel Circumference (m)\n"); p.add_run(f"    Circumference = ({inputs.get('wheel_diameter')} * π) / 1000\n"); self._add_highlighted_result(p, "    = ", results.get('wheel_circumference'), " m")
        p = doc.add_paragraph("  Calculate Wheel RPM\n"); p.add_run(f"    Wheel RPM = ({results.get('speed_mps'):.2f} m/s / {results.get('wheel_circumference'):.3f} m) * 60\n"); self._add_highlighted_result(p, "    = ", results.get('wheel_rpm'), " RPM")
        doc.add_heading("Step 2: Resistance Forces (kN)", level=3)
        p = doc.add_paragraph("  Rolling Resistance (kN)\n"); p.add_run(f"    = (A + B*Speed + C*Speed^2) * Weight * 9.81 / 1000\n"); self._add_highlighted_result(p, "    = ", results.get('rolling_resistance'), " kN")
        p = doc.add_paragraph("  Gradient Resistance (kN)\n"); p.add_run(f"    = Weight * 1000 * 9.81 * Slope(%) / 100000\n"); self._add_highlighted_result(p, "    = ", results.get('gradient_resistance'), " kN")
        p = doc.add_paragraph("  Curvature Resistance (kN)\n"); p.add_run(f"    = 0.4 * Weight * Curve(deg) * 9.81 / 1000\n"); self._add_highlighted_result(p, "    = ", results.get('curvature_resistance'), " kN")
        p = doc.add_paragraph("  Starting Resistance (kN)\n"); p.add_run(f"    = 6 * Weight * 9.81 / 1000\n"); self._add_highlighted_result(p, "    = ", results.get('starting_resistance'), " kN")
        p = doc.add_paragraph(); self._add_highlighted_result(p, "  TOTAL RESISTANCE = ", results.get('total_resistance'), " kN")
        doc.add_heading("Step 3: Torque Requirements", level=3)
        p = doc.add_paragraph("  Calculate Wheel Radius (m)\n"); p.add_run(f"    = {inputs.get('wheel_diameter')} / 2000\n"); self._add_highlighted_result(p, "    = ", results.get('wheel_radius'), " m")
        p = doc.add_paragraph("  Required Total Torque at Wheels (Nm)\n"); p.add_run(f"    = Total Resistance * 1000 * Wheel Radius\n"); self._add_highlighted_result(p, "    = ", results.get('required_total_torque'), " Nm")
        p = doc.add_paragraph("  Required Torque per Wheel (Nm)\n"); p.add_run(f"    = Total Torque / (Axles * 2)\n"); self._add_highlighted_result(p, "    = ", results.get('per_wheel_torque'), " Nm/wheel")
        p = doc.add_paragraph("  Required Torque per Axle (Nm)\n"); p.add_run(f"    = Total Torque / Axles\n"); self._add_highlighted_result(p, "    = ", results.get('per_axle_torque'), " Nm/axle")
        p = doc.add_paragraph("  Required Motor Torque (Nm) (per gearbox input)\n"); p.add_run(f"    = Torque per Axle / Axle Gear Box Ratio\n"); self._add_highlighted_result(p, "    = ", results.get('per_gearbox_input_torque'), " Nm")
        doc.add_heading("Step 4: Motor Speed", level=3)
        p = doc.add_paragraph("  Motor Speed (RPM) (gearbox input RPM)\n"); p.add_run(f"    = Wheel RPM * Axle Gear Box Ratio\n"); p.add_run(f"    = {results.get('wheel_rpm'):.2f} * {inputs.get('axle_gear_box_ratio')}\n"); self._add_highlighted_result(p, "    = ", results.get('gearbox_input_rpm'), " RPM")
        doc.add_heading("Step 5: Motor Displacement (New Formula)", level=3)
        p = doc.add_paragraph("  Convert Motor Torque (Nm) to (kg-cm)\n"); p.add_run(f"    = {results.get('per_gearbox_input_torque'):.2f} Nm * 10.1972\n"); self._add_highlighted_result(p, "    = ", results.get('per_gearbox_input_torque_kg_cm'), " kg-cm")
        p = doc.add_paragraph("  Convert Pressure (bar) to (kg/cm2)\n"); p.add_run(f"    = {inputs.get('pressure'):.2f} bar * 1.01972\n"); self._add_highlighted_result(p, "    = ", results.get('pressure_kg_cm2'), " kg/cm2")
        p = doc.add_paragraph("  Motor Displacement (cc/rev)\n"); p.add_run(f"    = (Torque (kg-cm) * 2 * 3.1416) / (Pressure (kg/cm2) * Mech Eff)\n"); p.add_run(f"    = ({results.get('per_gearbox_input_torque_kg_cm'):.2f} * 6.2832) / ({results.get('pressure_kg_cm2'):.2f} * {inputs.get('mech_eff_motor')/100.0})\n"); self._add_highlighted_result(p, "    = ", results.get('motor_displacement_cc'), " cc/rev")
        doc.add_heading("Step 6: Motor Flow Rate (New Formula)", level=3)
        p = doc.add_paragraph("  Per Motor Flow Rate (LPM)\n"); p.add_run(f"    = ((Motor Disp (cc) * Motor Speed) / Motor Vol Eff) / 1000\n"); p.add_run(f"    = (({results.get('motor_displacement_cc'):.2f} * {results.get('gearbox_input_rpm'):.2f}) / {inputs.get('vol_eff_motor')/100.0}) / 1000\n"); self._add_highlighted_result(p, "    = ", results.get('per_motor_flow_rate_lpm'), " LPM")
        doc.add_page_break()
        doc.add_heading("Step 7: Required Pump Displacement (Per Engine Gear)", level=2)
        table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells; hdr_cells[0].text = 'Engine Gear Ratio'; hdr_cells[1].text = 'max Vehicle RPM'; hdr_cells[2].text = 'Calculated Pump RPM'; hdr_cells[3].text = 'Required Pump Disp. (cc/rev)'
        pump_results_list = results.get('pump_results', [])
        for res in pump_results_list:
            row_cells = table.add_row().cells; row_cells[0].text = f"{res.get('engine_gear_ratio', 0):.2f}"; row_cells[1].text = f"{res.get('max_vehicle_rpm_input', 0):.0f}"; calc_pump_rpm = res.get('pump_rpm', 0); row_cells[2].text = f"{calc_pump_rpm:.2f}"; run = row_cells[3].paragraphs[0].add_run(f"{res.get('pump_disp_cc', 0):.2f}"); run.font.bold = True; run.font.color.rgb = RGBColor.from_string("0D47A1")
        doc.add_paragraph(); doc.add_heading("Step 7 (Details):", level=4)
        for res in pump_results_list:
            max_vehicle_rpm_input = res.get('max_vehicle_rpm_input', 0); actual_prop_rpm = res.get('actual_prop_rpm', 0); calc_pump_rpm = res.get('pump_rpm', 0); final_disp_cc = res.get('pump_disp_cc', 0); engine_gear = res.get('engine_gear_ratio', 1.0)
            doc.add_paragraph(f"  - For Engine Gear {engine_gear:.2f} @ {max_vehicle_rpm_input:.0f} RPM:", style='List Bullet')
            p = doc.add_paragraph("    Calculate Actual Prop RPM\n"); p.add_run(f"      = Max RPM / Engine Gear\n"); p.add_run(f"      = {max_vehicle_rpm_input:.0f} / {engine_gear}\n"); self._add_highlighted_result(p, f"      = ", actual_prop_rpm, " RPM")
            p = doc.add_paragraph("    Calculate Pump RPM\n"); p.add_run(f"      = Actual Prop RPM * PTO Gear box Ratio\n"); p.add_run(f"      = {actual_prop_rpm:.2f} * {inputs.get('pto_gear_ratio')}\n"); self._add_highlighted_result(p, f"      = ", calc_pump_rpm, " RPM")
            p = doc.add_paragraph("    Calculate Required Pump Displacement (cc/rev)\n"); p.add_run(f"      = (Per Motor Flow (LPM) * 1000) / (Calculated Pump RPM * Pump Vol Eff)\n"); p.add_run(f"      = ({results.get('per_motor_flow_rate_lpm'):.2f} * 1000) / ({calc_pump_rpm:.2f} * {inputs.get('vol_eff_pump')/100.0})\n"); self._add_highlighted_result(p, f"      = ", final_disp_cc, " cc/rev")
        
    # --- .docx Mode 2 ---
    def _create_mode2_docx(self, doc, inputs, results_wrapper):
        if docx is None: raise ImportError("python-docx library is required to generate .docx files.")
        doc.add_heading("Achievable Speed Calculation Report", level=0)
        raw = self.inputs_raw
        results = results_wrapper
        doc.add_heading("Input Parameters", level=2)
        vehicle_data = {
            "Wheel Dia (mm)": raw.get('wheel_diameter'), "PTO Gear box Ratio": raw.get('pto_gear_ratio'),
            "Axle Gear box Ratio": raw.get('axle_gear_box_ratio'), "Engine Gear Box Ratios": raw.get('engine_gear_ratio'),
            "max Vehicle RPM": raw.get('max_vehicle_rpm'),
        }
        hydraulic_data = {
            "Hydraulic Motor (Total)": raw.get('num_motors'), "Hydraulic Motor / axle": raw.get('per_axle_motor'),
            "Pump Displacement (cc)": raw.get('pump_disp_in'), "Pump Vol Eff (%)": raw.get('vol_eff_pump'),
            "max Limit of Pump RPM": raw.get('max_pump_rpm'), "Motor Displacement (cc)": raw.get('motor_disp_in'),
            "Motor Vol Eff (%)": raw.get('vol_eff_motor'), "max Motor RPM": raw.get('max_motor_rpm'),
        }
        self._populate_input_table(doc, "Vehicle Inputs", vehicle_data)
        self._populate_input_table(doc, "Hydraulic Inputs", hydraulic_data)
        doc.add_heading("Calculation Results (Summary)", level=2)
        table = doc.add_table(rows=1, cols=7); table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells; hdr_cells[0].text = 'Engine Gear Ratio'; hdr_cells[1].text = 'max Vehicle RPM'; hdr_cells[2].text = 'Calc. Pump RPM'; hdr_cells[3].text = 'Calc. Pump Flow (LPM)'; hdr_cells[4].text = 'Calc. Motor Speed (RPM)'; hdr_cells[5].text = 'Calc. Axle/Wheel Speed (RPM)'; hdr_cells[6].text = 'Achievable Speed (km/h)'
        speed_results_list = results.get('speed_results_list', [])
        for res in speed_results_list:
            row_cells = table.add_row().cells; row_cells[0].text = f"{res.get('engine_gear_ratio', 0):.2f}"; row_cells[1].text = f"{res.get('max_vehicle_rpm_input', 0):.0f}"
            calc_pump_rpm = res.get('pump_rpm', 0); run_pump_rpm = row_cells[2].paragraphs[0].add_run(f"{calc_pump_rpm:.2f}"); 
            if calc_pump_rpm > inputs.get('max_pump_rpm'): run_pump_rpm.font.color.rgb = RGBColor.from_string("FF0000"); run_pump_rpm.font.bold = True
            row_cells[3].text = f"{res.get('pump_flow_lpm', 0):.2f}"
            calc_motor_rpm = res.get('motor_speed_rpm', 0); run_motor_rpm = row_cells[4].paragraphs[0].add_run(f"{calc_motor_rpm:.2f}"); 
            if calc_motor_rpm > inputs.get('max_motor_rpm'): run_motor_rpm.font.color.rgb = RGBColor.from_string("FF0000"); run_motor_rpm.font.bold = True
            row_cells[5].text = f"{res.get('axle_shaft_rpm', 0):.2f}"
            run = row_cells[6].paragraphs[0].add_run(f"{res.get('achievable_speed_kph', 0):.2f}"); run.font.bold = True; run.font.color.rgb = RGBColor.from_string("0D47A1")
        p_note = doc.add_paragraph("Note: Values in "); run_red = p_note.add_run("red"); run_red.font.color.rgb = RGBColor.from_string("FF0000"); run_red.font.bold = True; p_note.add_run(f" exceed the specified limits (Max Motor: {inputs.get('max_motor_rpm'):.0f} RPM, Max Pump: {inputs.get('max_pump_rpm'):.0f} RPM).")
        doc.add_page_break()
        doc.add_heading("Step-by-Step Calculation Details", level=2)
        if not speed_results_list: doc.add_paragraph(" (No speed results were calculated)")
        for res in speed_results_list:
            max_vehicle_rpm_input = res.get('max_vehicle_rpm_input', 0); actual_prop_rpm = res.get('actual_prop_rpm', 0); pump_rpm = res.get('pump_rpm', 0); pump_flow_lpm = res.get('pump_flow_lpm', 0); motor_speed_rpm = res.get('motor_speed_rpm', 0); axle_shaft_rpm = res.get('axle_shaft_rpm', 0); wheel_circumference = res.get('wheel_circumference', 0); achievable_speed_kph = res.get('achievable_speed_kph', 0); motor_disp_lpm = res.get('motor_disp_lpm', 0); engine_gear = res.get('engine_gear_ratio', 1.0)
            doc.add_heading(f"--- For Engine Gear {engine_gear:.2f} @ {max_vehicle_rpm_input:.0f} RPM ---", level=3)
            p = doc.add_paragraph("Calculate Actual Prop RPM\n"); p.add_run(f"  = Max RPM / Engine Gear\n"); p.add_run(f"  = {max_vehicle_rpm_input:.0f} / {engine_gear}\n"); self._add_highlighted_result(p, f"  = ", actual_prop_rpm, " RPM")
            p = doc.add_paragraph("Calculate Pump RPM\n"); p.add_run(f"  = Actual Prop RPM * PTO Gear box Ratio\n"); p.add_run(f"  = {actual_prop_rpm:.2f} * {inputs.get('pto_gear_ratio')}\n"); self._add_highlighted_result(p, f"  = ", pump_rpm, " RPM")
            p_pump_limit = doc.add_paragraph(); p_pump_limit.add_run(f"  (Max Pump RPM Limit: {inputs.get('max_pump_rpm'):.0f} RPM)"); 
            if pump_rpm > inputs.get('max_pump_rpm'): run = p_pump_limit.add_run("    ⚠️ WARNING: Exceeds Limit!"); run.font.color.rgb = RGBColor.from_string("FF0000"); run.font.bold = True
            p = doc.add_paragraph("Calculate Pump Flow (LPM)\n"); p.add_run(f"  = (Pump Disp (cc) * Pump RPM * Pump Vol Eff) / 1000\n"); p.add_run(f"  = ({inputs.get('pump_disp_in')} * {pump_rpm:.2f} * {inputs.get('vol_eff_pump')/100.0}) / 1000\n"); self._add_highlighted_result(p, f"  = ", pump_flow_lpm, " LPM")
            p = doc.add_paragraph("Calculate Motor Speed (RPM)\n"); p.add_run(f"  = (Pump Flow (LPM) / Motor Disp (L/rev)) * Motor Vol Eff\n"); p.add_run(f"  = ({pump_flow_lpm:.2f} / {motor_disp_lpm:.3f}) * {inputs.get('vol_eff_motor')/100.0}\n"); self._add_highlighted_result(p, f"  = ", motor_speed_rpm, " RPM")
            p_motor_limit = doc.add_paragraph(); p_motor_limit.add_run(f"  (Max Motor RPM Limit: {inputs.get('max_motor_rpm'):.0f} RPM)"); 
            if motor_speed_rpm > inputs.get('max_motor_rpm'): run = p_motor_limit.add_run("    ⚠️ WARNING: Exceeds Limit!"); run.font.color.rgb = RGBColor.from_string("FF0000"); run.font.bold = True
            p = doc.add_paragraph("Calculate Axle/Wheel Speed (RPM)\n"); p.add_run(f"  = Motor Speed / Axle Gear box Ratio\n"); p.add_run(f"  = {motor_speed_rpm:.2f} / {inputs.get('axle_gear_box_ratio')}\n"); self._add_highlighted_result(p, f"  = ", axle_shaft_rpm, " RPM")
            doc.add_heading("Final Speed Calculation", level=4)
            p = doc.add_paragraph(f"  Wheel Circumference = {wheel_circumference:.3f} m\n"); p.add_run(f"  Speed (m/s) = (Axle/Wheel Speed * Wheel Circumference) / 60\n"); p.add_run(f"  Speed (m/s) = ({axle_shaft_rpm:.2f} * {wheel_circumference:.3f}) / 60 = {achievable_speed_kph / 3.6:.2f} m/s\n"); p.add_run(f"  Speed (km/h) = Speed (m/s) * 3.6\n"); self._add_highlighted_result(p, f"  = {achievable_speed_kph / 3.6:.2f} * 3.6 = ", achievable_speed_kph, " km/h")
            doc.add_page_break()