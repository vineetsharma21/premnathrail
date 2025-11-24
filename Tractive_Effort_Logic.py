import math
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Core Calculation Logic ---
def perform_te_calculations(inputs):
    """
    Performs all tractive effort calculations and returns the results as a dictionary.
    Assumes all inputs in the 'inputs' dict are pre-validated (are floats or correct strings).
    """
    # Extract validated inputs
    load = inputs['load']
    loco_weight = inputs['loco_weight']
    gradient_input = inputs['gradient']
    curvature_input = inputs['curvature']
    speed = inputs['speed']
    mode = inputs['mode']
    grad_type = inputs['grad_type']
    curvature_unit = inputs['curvature_unit']
    
    results = {}
    
    # Gradient Resistance
    if grad_type == "Degree":
        gradient_resistance_per_ton = math.tan(math.radians(gradient_input)) * 1000 if gradient_input != 0 else 0
    else: # "1 in G"
        gradient_resistance_per_ton = 1000 / gradient_input if gradient_input != 0 else 0
        
    # Curvature Resistance
    if curvature_unit == "Radius(m)":
        curvature_resistance_per_ton = 700 / curvature_input if curvature_input != 0 else 0
    else: # "Degree"
        curvature_resistance_per_ton = curvature_input
        
    # Mode-based values
    if mode == "Start":
        wagon_rolling_resistance, loco_rolling_resistance, speed_for_power = 4.0, 6.0, 1.0
    else: # "Running"
        wagon_rolling_resistance, loco_rolling_resistance, speed_for_power = 1.3505, 2.913, speed
        
    # Calculate resistances
    results['T1'] = load * wagon_rolling_resistance
    results['T2'] = loco_weight * loco_rolling_resistance
    results['T3'] = (load + loco_weight) * gradient_resistance_per_ton
    results['T4'] = (load + loco_weight) * curvature_resistance_per_ton
    
    # Calculate final results
    results['te'] = results['T1'] + results['T2'] + results['T3'] + results['T4']
    results['power'] = (results['te'] * speed_for_power) / 270
    results['ohe_current'] = (results['power'] * 735.5) / (22500 * 0.84 * 0.8)
    
    return results

def format_te_report_text(inputs, results):
    """Formats the calculation inputs and results into a readable text report."""
    return (
        f"# Tractive Effort Calculation Report\n\n"
        f"--- 1. Inputs ---\n"
        f"• Shunting Load: {inputs['load']} tons\n"
        f"• GBW of Vehicle: {inputs['loco_weight']} tons\n"
        f"• Gradient: {inputs['gradient']} ({inputs['grad_type']})\n"
        f"• Curvature: {inputs['curvature']} ({inputs['curvature_unit']})\n"
        f"• Speed: {inputs['speed']} km/h\n"
        f"• Mode: {inputs['mode']}\n\n"
        
        f"--- 2. Calculation Results ---\n"
        f"Summary of Results:\n"
        f"  • Tractive Effort (TE): {results['te']:.2f} kg  ({results['te']/1000:.3f} tons)\n"
        f"  • Rail Horsepower: {results['power']:.2f} HP\n"
        f"  • OHE Current: {results['ohe_current']:.2f} A\n\n"
        f"Resistance Components:\n"
        f"  • T1 (Wagon Rolling Resistance): {results['T1']:.2f} kg\n"
        f"  • T2 (Loco Rolling Resistance): {results['T2']:.2f} kg\n"
        f"  • T3 (Gradient Resistance): {results['T3']:.2f} kg\n"
        f"  • T4 (Curvature Resistance): {results['T4']:.2f} kg\n"
    )

def create_te_report_docx(inputs, results):
    """
    Creates a .docx report for Tractive Effort and returns it as an in-memory stream.
    Adapted from the tkinter file.
    """
    doc = Document()
    doc.add_heading("Tractive Effort Calculation Report", 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    doc.add_heading('1. Inputs', level=1)
    p = doc.add_paragraph()
    p.add_run(f"• Shunting Load: {inputs['load']} tons\n")
    p.add_run(f"• GBW of Vehicle: {inputs['loco_weight']} tons\n")
    p.add_run(f"• Gradient: {inputs['gradient']} ({inputs['grad_type']})\n")
    p.add_run(f"• Curvature: {inputs['curvature']} ({inputs['curvature_unit']})\n")
    p.add_run(f"• Speed: {inputs['speed']} km/h\n")
    p.add_run(f"• Mode: {inputs['mode']}\n")
    
    doc.add_heading('2. Calculation Results', level=1)
    p = doc.add_paragraph()
    run = p.add_run("Summary of Results:\n")
    run.bold = True
    p.add_run(f"  • Tractive Effort (TE): {results['te']:.2f} kg  ({results['te']/1000:.3f} tons)\n")
    p.add_run(f"  • Rail Horsepower: {results['power']:.2f} HP\n")
    p.add_run(f"  • OHE Current: {results['ohe_current']:.2f} A\n")
    
    p = doc.add_paragraph()
    run = p.add_run("\nResistance Components:\n")
    run.bold = True
    p.add_run(f"  • T1 (Wagon Rolling Resistance): {results['T1']:.2f} kg\n")
    p.add_run(f"  • T2 (Loco Rolling Resistance): {results['T2']:.2f} kg\n")
    p.add_run(f"  • T3 (Gradient Resistance): {results['T3']:.2f} kg\n")
    p.add_run(f"  • T4 (Curvature Resistance): {results['T4']:.2f} kg\n")

    try:
        # Save to in-memory stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        raise RuntimeError(f"An error occurred while creating the docx report: {e}")